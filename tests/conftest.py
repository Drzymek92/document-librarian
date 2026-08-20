import sys
from pathlib import Path

# Put the project root first on sys.path so `import scripts.*` resolves to the
# local package and not an unrelated `scripts` package in site-packages.
PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

import hashlib
import math
import re

import pytest

# --- Deterministic stand-in for a real embedding model ----------------------
# Hashing bag-of-words: each token lands in a fixed bucket, so two texts sharing
# vocabulary get a high cosine and two that share none get ~0. Enough to exercise
# ranking, fusion and the model lock without a network call or an API key, and
# deterministic so a failure is a real regression rather than model drift.
TOY_DIM = 64


def toy_embed(text: str) -> list[float]:
    vec = [0.0] * TOY_DIM
    for token in re.findall(r"[a-z0-9]+", text.lower()):
        bucket = int(hashlib.sha1(token.encode("utf-8")).hexdigest()[:8], 16) % TOY_DIM
        vec[bucket] += 1.0
    norm = math.sqrt(sum(v * v for v in vec))
    return [v / norm for v in vec] if norm else vec


@pytest.fixture
def toy_embedder(monkeypatch):
    """Point the embedding layer at toy_embed for the duration of a test."""
    from scripts import embed

    monkeypatch.setattr(
        embed, "embed_texts", lambda texts, model: [toy_embed(t) for t in texts]
    )
    monkeypatch.setattr(embed, "embed_query", lambda text, model: toy_embed(text))
    return toy_embed


@pytest.fixture(scope="session")
def sample_dir(tmp_path_factory):
    d = tmp_path_factory.mktemp("samples")

    # CSV
    (d / "data.csv").write_text(
        "name,amount\nAlice,100\nBob,250\n", encoding="utf-8"
    )

    # XLSX
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Q1"
    ws.append(["product", "units"])
    ws.append(["widget", 42])
    ws.append(["gadget", 7])
    wb.save(d / "sheet.xlsx")

    # PDF (two pages)
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(d / "doc.pdf"))
    c.drawString(72, 720, "Quarterly report the quick brown fox")
    c.showPage()
    c.drawString(72, 720, "Second page lazy dog sleeps")
    c.showPage()
    c.save()

    # DOCX
    import docx

    document = docx.Document()
    document.add_heading("Title Memo", level=0)
    document.add_paragraph("The quick brown fox.")
    document.add_paragraph("Lazy dog sleeps.")
    document.save(d / "memo.docx")

    # DOCX with a table (content lives only in table cells)
    doc_t = docx.Document()
    doc_t.add_paragraph("Inventory snapshot.")
    table = doc_t.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "product"
    table.rows[0].cells[1].text = "units"
    table.rows[1].cells[0].text = "sprocket"
    table.rows[1].cells[1].text = "314"
    doc_t.save(d / "table_memo.docx")

    # Semicolon-delimited, cp1252-encoded CSV (locale export)
    (d / "locale.csv").write_bytes(
        "name;city\nRenée;Montréal\nJosé;Bogotá\n".encode("cp1252")
    )

    # Markdown with headings
    (d / "guide.md").write_text(
        "# Overview\n\nIntro text here.\n\n"
        "## Locale Setup\n\nConfigure the locale carefully.\n\n"
        "## Troubleshooting\n\nCheck the encoding first.\n",
        encoding="utf-8",
    )

    # Image-only PDF (blank page → no text layer, triggers OCR fallback)
    from reportlab.pdfgen import canvas as _canvas

    cb = _canvas.Canvas(str(d / "scanned.pdf"))
    cb.showPage()
    cb.save()

    # JSON (array of nested records, mimics an FTS job export)
    (d / "feedback.json").write_text(
        '[{"id": 1, "locale": "fr-CA", "note": {"text": "needs review"}},'
        ' {"id": 2, "locale": "es-CO", "note": {"text": "approved"}}]',
        encoding="utf-8",
    )

    # JSONL (one record per line)
    (d / "events.jsonl").write_text(
        '{"event": "ingested", "count": 42}\n{"event": "transcribed", "count": 7}\n',
        encoding="utf-8",
    )

    # TSV (tab-delimited — exercises the CSV delimiter sniffer)
    (d / "prices.tsv").write_text(
        "sku\tprice\nA-100\t9.99\nB-200\t14.50\n", encoding="utf-8"
    )

    # HTML (with script/style that must be stripped)
    (d / "page.html").write_text(
        "<html><head><title>Locale Report</title>"
        "<style>body{color:red}</style></head>"
        "<body><h1>Summary</h1><p>The transcription quality improved.</p>"
        "<script>console.log('noise')</script></body></html>",
        encoding="utf-8",
    )

    # Plain text
    (d / "notes.txt").write_text(
        "Locale feedback notes. The transcription quality needs review.\n",
        encoding="utf-8",
    )

    # Image
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (300, 80), "white")
    ImageDraw.Draw(img).text((10, 30), "Receipt 99.50", fill="black")
    img.save(d / "scan.png")

    return d
